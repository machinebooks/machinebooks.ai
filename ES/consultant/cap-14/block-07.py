# Extraído de: LibroConsultor/cap-14-reporting.md
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_one_pager(
    proyecto: ProyectoReporting,
    resumen_ejecutivo: str,
    top_hallazgos: list[Hallazgo],
    output_path: str,
    plantilla_docx: str
) -> str:
    """Genera resumen ejecutivo de una página."""

    doc = Document(plantilla_docx)

    # Título
    titulo = doc.add_heading(proyecto.nombre_proyecto, level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"{proyecto.cliente} | {proyecto.tipo.upper()} | "
        f"{proyecto.fecha_fin}"
    ).font.size = Pt(10)

    # Resumen en 3 frases
    doc.add_heading("Resumen", level=2)
    # Usar solo el primer párrafo del resumen ejecutivo
    primer_parrafo = resumen_ejecutivo.split("\n\n")[0]
    doc.add_paragraph(primer_parrafo)

    # Hallazgos críticos en tabla compacta
    doc.add_heading("Hallazgos principales", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Hallazgo"
    hdr[1].text = "Severidad"
    hdr[2].text = "Acción requerida"

    for h in top_hallazgos[:5]:
        row = table.add_row().cells
        row[0].text = h.titulo
        row[1].text = h.severidad.value.upper()
        row[2].text = h.recomendacion[:80]  # Truncar

    # Próximos pasos
    doc.add_heading("Próximos pasos", level=2)
    doc.add_paragraph(
        "Sesión de validación de hallazgos con los equipos "
        "técnicos en los próximos 5 días hábiles."
    )

    doc.save(output_path)
    return output_path
